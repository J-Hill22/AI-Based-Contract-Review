import pandas as pd
import re
import docx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Main Program
def annotate_contract(xls_in, contract_in, save_path):
    # Read FAR Matrix and clean it
    xls_matrix_path = xls_in
    df_clean = read_far_matrix(xls_matrix_path)
    xls_clauses = df_clean['Clause'].tolist()
    acceptance_status = df_clean['Acceptance Status*'].tolist()

    # Read Contract Text
    contract_text = read_docx(contract_in)

    # Flag Clauses in Contract Text
    flagged_clauses = flag_clauses(contract_text, xls_clauses, acceptance_status)

    # Create New Annotated Document
    new_doc = Document()
    original_doc = Document(contract_in)
    for paragraph in original_doc.paragraphs:
        annotate_paragraph(paragraph, flagged_clauses, new_doc)
    
    # Save Annotated Document
    new_doc.save(save_path)

# Function to read FAR Matrix from Excel and return clean DataFrame
def read_far_matrix(excel_path):
    df = pd.read_excel(excel_path)
    df_clean = df.dropna(subset=['Clause ', 'Acceptance Status*'])
    df_clean = df_clean[['Clause ', 'Acceptance Status*']]
    df_clean.rename(columns={'Clause ': 'Clause'}, inplace=True)
    return df_clean

# Function to read text from a .docx file
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to flag clauses in contract text
def flag_clauses(contract_text, xls_clauses, acceptance_status):
    flagged_clauses = {}
    for clause, status in zip(xls_clauses, acceptance_status):
        if re.search(re.escape(clause), contract_text, re.IGNORECASE):
            flagged_clauses[clause] = status
    return flagged_clauses

# Function to annotate a paragraph based on flagged clauses and their acceptance status
def annotate_paragraph(paragraph, flagged_clauses, new_doc):
    new_paragraph = new_doc.add_paragraph()
    for run in paragraph.runs:
        new_run = new_paragraph.add_run(run.text, run.style.name)
        new_run.font.highlight_color = run.font.highlight_color if run.font.highlight_color else None
        for clause, status in flagged_clauses.items():
            if re.search(re.escape(clause), run.text, re.IGNORECASE):
                if status == 'OK':
                    new_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN # Green for OK
                elif status == 'C':
                    new_run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE  # Turqoise for Conditional (would be blue but its hard to read)
                elif status == 'REMOVE':
                    new_run.font.highlight_color = WD_COLOR_INDEX.RED  # Red for Remove